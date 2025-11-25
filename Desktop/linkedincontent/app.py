import streamlit as st
import requests
import pdfplumber
import io
import time
import textwrap
import pandas as pd
import matplotlib.pyplot as plt
import base64
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- 全局配置 ---
st.set_page_config(page_title="Pro Research Agent (Copy Ready)", layout="wide", page_icon="💎")

# 绘图配置 (解决中文)
plt.style.use('ggplot')
plt.rcParams['font.family'] = 'sans-serif'
plt.rcParams['font.sans-serif'] = ['SimHei', 'Arial', 'Microsoft YaHei'] 
plt.rcParams['axes.unicode_minus'] = False

# --- 状态管理 ---
if 'history' not in st.session_state:
    st.session_state['history'] = []
if 'current_report' not in st.session_state:
    st.session_state['current_report'] = None

# --- 核心函数 ---

def extract_pages_from_pdf(uploaded_file):
    pages_content = []
    with pdfplumber.open(uploaded_file) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                pages_content.append(text)
    return pages_content

def call_ai_api(api_key, base_url, model_name, messages, temperature=0.1):
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model_name, "messages": messages, "temperature": temperature, "stream": False}
    try:
        response = requests.post(base_url, headers=headers, json=payload, timeout=300)
        if response.status_code == 200:
            return response.json()['choices'][0]['message']['content']
        return None 
    except Exception:
        return None

def create_table_image_bytes(markdown_table_lines):
    """
    生成表格图片，返回 BytesIO 对象
    """
    try:
        clean_rows = []
        for line in markdown_table_lines:
            content = line.strip()
            if '|' in content:
                clean_check = content.replace('|', '').replace('-', '').replace(':', '').strip()
                if clean_check: 
                    clean_rows.append(content)
        
        if len(clean_rows) < 2: return None
        
        data_matrix = []
        max_cols = 0
        for line in clean_rows:
            line_pure = line.strip().strip('|')
            cells = [c.strip() for c in line_pure.split('|')]
            data_matrix.append(cells)
            if len(cells) > max_cols: max_cols = len(cells)

        final_data = []
        for row in data_matrix:
            if len(row) < max_cols:
                row += [""] * (max_cols - len(row))
            final_data.append(row[:max_cols])
            
        if not final_data: return None

        headers = final_data[0]
        body = final_data[1:]
        if not body: body = [[""] * len(headers)]
        
        df = pd.DataFrame(body, columns=headers)

        # 绘图逻辑
        row_heights = []
        col_width_chars = 20
        for row in body:
            max_lines = 1
            for cell in row:
                lines = len(textwrap.wrap(str(cell), width=col_width_chars))
                if lines > max_lines: max_lines = lines
            row_heights.append(max_lines)
            
        base_h = 0.5
        total_h = 0.8 + sum([rh * base_h for rh in row_heights])
        total_w = min(len(headers) * 3.0, 12)

        fig, ax = plt.subplots(figsize=(total_w, total_h))
        ax.axis('off')
        
        table = ax.table(cellText=df.values, colLabels=df.columns, loc='center', cellLoc='left')
        table.auto_set_font_size(False)
        table.set_fontsize(11)
        
        cells = table.get_celld()
        for (row, col), cell in cells.items():
            cell.set_edgecolor('#bfbfbf')
            cell.set_linewidth(1)
            cell.set_text_props(position=(0.02, cell.get_text_props()['position'][1])) 
            
            if row == 0:
                cell.set_height(0.8 / total_h)
                cell.set_facecolor('#2c3e50')
                cell.set_text_props(color='white', weight='bold', ha='center', fontsize=12)
            else:
                rh_mult = row_heights[row-1]
                cell.set_height((rh_mult * base_h) / total_h)
                cell.set_facecolor('#f8f9fa' if row % 2 else '#ffffff')
                cell.set_text_props(color='#333333', wrap=True, ha='left', va='center')

        img_buffer = io.BytesIO()
        plt.savefig(img_buffer, format='png', bbox_inches='tight', dpi=150, pad_inches=0.1)
        plt.close(fig)
        img_buffer.seek(0)
        return img_buffer
    except Exception:
        return None

def parse_blocks(text_content):
    """
    将文本解析为 Block 列表，同时处理图片 Base64 编码以便 HTML 使用
    """
    lines = text_content.split('\n')
    blocks = [] 
    
    current_text = []
    table_buffer = []
    inside_table = False
    
    for line in lines:
        stripped = line.strip()
        is_table_row = '|' in stripped and len(stripped) > 3
        
        if is_table_row:
            if not inside_table:
                if current_text:
                    blocks.append({'type': 'text', 'content': "\n".join(current_text)})
                    current_text = []
                inside_table = True
            table_buffer.append(stripped)
        else:
            if inside_table:
                # 生成表格图片
                img_bytes = create_table_image_bytes(table_buffer)
                if img_bytes:
                    # 关键：转为 Base64 字符串
                    b64_str = base64.b64encode(img_bytes.getvalue()).decode()
                    blocks.append({
                        'type': 'image', 
                        'bytes': img_bytes,   # 给 Word 用
                        'base64': b64_str     # 给 HTML 用
                    })
                else:
                    current_text.extend(table_buffer)
                
                inside_table = False
                table_buffer = []
                
            current_text.append(line)
            
    if inside_table and table_buffer:
        img_bytes = create_table_image_bytes(table_buffer)
        if img_bytes:
            b64_str = base64.b64encode(img_bytes.getvalue()).decode()
            blocks.append({'type': 'image', 'bytes': img_bytes, 'base64': b64_str})
        else:
            current_text.extend(table_buffer)
            
    if current_text:
        blocks.append({'type': 'text', 'content': "\n".join(current_text)})
        
    return blocks

def generate_word(blocks):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    
    doc.add_heading('Research Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("_" * 50)

    for block in blocks:
        if block['type'] == 'text':
            for line in block['content'].split('\n'):
                s = line.strip()
                if not s: continue
                if s.startswith('# '): doc.add_heading(s[2:], 1)
                elif s.startswith('## '): doc.add_heading(s[3:], 2)
                elif s.startswith('### '): doc.add_heading(s[4:], 3)
                elif s.startswith('- '): doc.add_paragraph(s[2:], style='List Bullet')
                else: doc.add_paragraph(s)
        elif block['type'] == 'image':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                block['bytes'].seek(0) # 关键：重置指针
                run.add_picture(block['bytes'], width=Inches(6.0))
            except: pass

    bio = io.BytesIO()
    doc.save(bio)
    return bio

def generate_copyable_html(blocks):
    """
    生成一个包含嵌入式 Base64 图片的纯 HTML 字符串。
    这种格式可以被直接复制到 Email、Word、Notion 中而图片不丢失。
    """
    html = """
    <div id="copy-target" style="font-family: Arial, sans-serif; line-height: 1.6; color: #333; background-color: white; padding: 20px;">
    """
    
    for block in blocks:
        if block['type'] == 'text':
            # 简单的 Markdown 转 HTML
            text = block['content']
            # 转义 HTML 字符
            text = text.replace("<", "&lt;").replace(">", "&gt;")
            
            lines = text.split('\n')
            for line in lines:
                s = line.strip()
                if not s: continue
                if s.startswith('### '): html += f"<h3 style='color:#2c3e50; margin-top:15px;'>{s[4:]}</h3>"
                elif s.startswith('## '): html += f"<h2 style='color:#2c3e50; border-bottom:1px solid #eee; padding-bottom:5px;'>{s[3:]}</h2>"
                elif s.startswith('# '): html += f"<h1 style='color:#2c3e50;'>{s[2:]}</h1>"
                elif s.startswith('- ') or s.startswith('* '): html += f"<li style='margin-left:20px;'>{s[2:]}</li>"
                else: html += f"<p style='margin-bottom:10px;'>{s}</p>"
                
        elif block['type'] == 'image':
            # 使用 Base64 直接嵌入图片
            b64 = block['base64']
            html += f"""
            <div style="text-align: center; margin: 20px 0;">
                <img src="data:image/png;base64,{b64}" style="max-width: 100%; border: 1px solid #ddd; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
            </div>
            """
            
    html += "</div>"
    return html

# --- UI ---
with st.sidebar:
    st.title("🗃️ 历史记录")
    if st.session_state['history']:
        for i, item in enumerate(reversed(st.session_state['history'])):
            if st.button(f"Load: {item['time']}", key=f"hist_{i}"):
                st.session_state['current_report'] = item
                st.rerun()
    st.divider()
    api_key = st.text_input("API Key", value="sk-3UIO8MwTblfyQuEZz2WUCzQOuK4QwwIPALVcNxFFNUxJayu7", type="password")
    model_name = st.selectbox("Model", ["gemini-3-pro", "gpt-4o", "qwen-max"])

st.title("💎 Pro Research Agent (Visual Copy Ready)")

uploaded_file = st.file_uploader("上传 PDF", type=['pdf'])

if uploaded_file and st.button("🔥 开始"):
    api_url = "https://api.nuwaapi.com/v1/chat/completions"
    
    with st.spinner("📖 正在 OCR 识别..."):
        pages = extract_pages_from_pdf(uploaded_file)

    full_text_parts = []
    progress = st.progress(0)
    
    for i, p in enumerate(pages):
        # 强制 AI 输出 Markdown 表格
        prompt = "You are an OCR engine. Output exact text. Detect tables and format them as Markdown Tables (| col |...)."
        msg = [{"role": "user", "content": f"{prompt}\n\n{p}"}]
        res = call_ai_api(api_key, api_url, model_name, msg)
        full_text_parts.append(res if res else p)
        progress.progress((i+1)/len(pages))

    full_text = "\n\n".join(full_text_parts)

    with st.spinner("🎨 生成可视化表格与 HTML..."):
        # 解析文本，生成图片对象
        blocks = parse_blocks(full_text)
        # 生成可复制的 HTML 代码
        html_content = generate_copyable_html(blocks)
        # 生成 Word
        word_data = generate_word(blocks)

    with st.spinner("🧠 撰写社媒..."):
        msg_s = [{"role": "user", "content": f"Act as Lead Analyst. Write social media posts.\n\n{full_text[:8000]}"}]
        social = call_ai_api(api_key, api_url, model_name, msg_s, temperature=0.7)

    report = {
        "time": datetime.now().strftime("%H:%M"),
        "filename": uploaded_file.name,
        "blocks": blocks,
        "html": html_content,
        "word": word_data.getvalue(),
        "social": social
    }
    st.session_state['current_report'] = report
    st.session_state['history'].append(report)
    st.rerun()

# --- 结果展示区 ---
curr = st.session_state['current_report']

if curr:
    st.divider()
    col1, col2 = st.columns([7, 3])
    
    with col1:
        st.subheader("📋 一键复制区 (完美保留表格图片)")
        
        # 下载 Word (以防万一)
        st.download_button("📂 下载 Word", curr['word'], "Report.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        
        st.markdown("""
        <div style="background-color: #e8f0fe; padding: 15px; border-radius: 5px; border-left: 5px solid #4285f4; margin-bottom: 20px;">
            <strong>🚀 如何复制：</strong> 
            <br>下面显示的是一个完整的 HTML 页面。
            <br>请在下方白色区域内 <strong>全选 (Ctrl + A)</strong> -> <strong>复制 (Ctrl + C)</strong>。
            <br>然后直接粘贴到 Word、邮件或微信中，<strong>表格图片会完美保留！</strong>
        </div>
        """, unsafe_allow_html=True)

        # --- 核心：渲染包含 Base64 图片的 HTML ---
        # 这是一个 iframe 或者 div，里面的图片是内嵌的，不是链接
        html_view = curr['html']
        
        # 我们使用一个带边框的容器来包裹这个 HTML，模拟一张“纸”
        st.markdown(f"""
        <div style="border: 1px solid #ddd; padding: 40px; border-radius: 2px; background: white; box-shadow: 0 4px 6px rgba(0,0,0,0.1);">
            {html_view}
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.subheader("🔥 社媒文案")
        st.text_area("Social Media", value=curr['social'], height=800)

elif not uploaded_file:
    st.info("请上传文件。系统将生成【内嵌 Base64 图片】的 HTML 视图，支持直接复制粘贴。")
